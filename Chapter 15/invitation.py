from docx import Document
doc=Document('Chapter 15/invitations.docx')
spacing=len('It would be a pleasure to have the company of')
print(spacing)
guests=open('Chapter 15/guests.txt')
guests=guests.read()
guests=guests.split("\n")
# print(guests)
styles=doc.styles
for i in guests:
    doc.add_paragraph("It would be a pleasure to have the company of").style='Intense Quote'
    doc.add_paragraph((spacing//2-len(i)//2)*" "+i+(spacing//2-len(i)//2)*" ")
    doc.add_paragraph( "at 11010 Memory lane in the Evening of").style='Intense Quote'
    doc.add_paragraph("April 1st")
    doc.add_paragraph("at 7 o'clock").style='Intense Quote'
    doc.add_page_break()
    
doc.save('Chapter 15/invitations.docx')

